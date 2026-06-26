"""Create the styled DOCX template used by the style-preservation example."""

from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


def shade_cell(cell, fill: str) -> None:
    """Apply a background fill color to a table cell."""

    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    tc_pr.append(shading)


def set_cell_text(cell, text: str, *, bold: bool = False, color: RGBColor | None = None) -> None:
    """Set cell text while preserving a single editable run."""

    cell.text = text
    run = cell.paragraphs[0].runs[0]
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def add_section_heading(document: Document, text: str) -> None:
    """Add a blue section heading with a thin rule underneath."""

    heading = document.add_paragraph()
    heading.paragraph_format.space_before = Pt(12)
    heading.paragraph_format.space_after = Pt(4)
    run = heading.add_run(text)
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(31, 78, 121)

    rule = document.add_paragraph()
    rule.paragraph_format.space_after = Pt(6)
    rule_run = rule.add_run(" ")
    rule_run.font.size = Pt(1)
    paragraph_border = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "9CC2E5")
    paragraph_border.append(bottom)
    rule._p.get_or_add_pPr().append(paragraph_border)


def main() -> None:
    out = Path("examples/templates/styled-status-report.docx")
    out.parent.mkdir(parents=True, exist_ok=True)

    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    normal = document.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10)

    header = section.header
    header_table = header.add_table(rows=1, cols=2, width=Inches(6.8))
    header_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    left, right = header_table.rows[0].cells
    shade_cell(left, "1F4E79")
    shade_cell(right, "1F4E79")
    set_cell_text(left, "STENCIL STATUS", bold=True, color=RGBColor(255, 255, 255))
    set_cell_text(right, "{{ report_date }}", bold=True, color=RGBColor(255, 255, 255))
    right.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run("Generated from a DOCX template with Stencil")
    footer_run.font.size = Pt(8)
    footer_run.font.color.rgb = RGBColor(127, 127, 127)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(0)
    title_run = title.add_run("Project Status Report")
    title_run.bold = True
    title_run.font.size = Pt(24)
    title_run.font.color.rgb = RGBColor(31, 78, 121)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(12)
    subtitle_run = subtitle.add_run("{{ project.name }}")
    subtitle_run.italic = True
    subtitle_run.font.size = Pt(13)
    subtitle_run.font.color.rgb = RGBColor(89, 89, 89)

    add_section_heading(document, "Executive Summary")

    summary = document.add_paragraph()
    summary.add_run("{{ summary }}")

    add_section_heading(document, "Health Snapshot")

    meta = document.add_table(rows=3, cols=2)
    meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta.style = "Table Grid"
    rows = [
        ("Owner", "{{ owner.name }}"),
        ("Status", "{{ status }}"),
        ("Reporting date", "{{ report_date }}"),
    ]
    for index, (label, value) in enumerate(rows):
        label_cell, value_cell = meta.rows[index].cells
        set_cell_text(label_cell, label, bold=True)
        set_cell_text(value_cell, value)
        shade_cell(label_cell, "EAF2F8")
        shade_cell(value_cell, "FFFFFF")
        if label == "Status":
            value_cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 112, 192)

    add_section_heading(document, "Milestones")

    table = document.add_table(rows=4, cols=3)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ["Milestone", "Owner", "Due"]
    for index, text in enumerate(headers):
        cell = table.rows[0].cells[index]
        set_cell_text(cell, text, bold=True, color=RGBColor(255, 255, 255))
        shade_cell(cell, "1F4E79")

    milestone_rows = [
        ("{{ milestones[0].name }}", "{{ milestones[0].owner }}", "{{ milestones[0].due }}"),
        ("{{ milestones[1].name }}", "{{ milestones[1].owner }}", "{{ milestones[1].due }}"),
        ("{{ milestones[2].name }}", "{{ milestones[2].owner }}", "{{ milestones[2].due }}"),
    ]
    for row_index, row_values in enumerate(milestone_rows, start=1):
        fill = "F7FBFD" if row_index % 2 else "FFFFFF"
        for cell_index, text in enumerate(row_values):
            cell = table.rows[row_index].cells[cell_index]
            set_cell_text(cell, text)
            shade_cell(cell, fill)

    add_section_heading(document, "Risk Note")

    note = document.add_paragraph()
    note_run = note.add_run("{% if risk_note %}Risk note: {{ risk_note }}{% endif %}")
    note_run.bold = True
    note_run.font.color.rgb = RGBColor(192, 0, 0)

    document.save(out)
    print(out)


if __name__ == "__main__":
    main()
